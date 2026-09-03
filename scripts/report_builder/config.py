# -*- coding: utf-8 -*-
"""
Configuration and Helper Functions for Report Generator
"""

import os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
REPORT_DOCX = os.path.join(BASE_DIR, "report", "Baocao.docx")
BACKUP_DOCX = os.path.join(BASE_DIR, "report", "Baocao_backup.docx")

# Màu sắc chuẩn mực học thuật hiện đại
COLOR_NAVY = RGBColor(15, 44, 89)         # #0F2C59 - Tiêu đề cấp 1
COLOR_DARK_BLUE = RGBColor(30, 58, 138)   # #1E3A8A - Tiêu đề cấp 2
COLOR_DARK_GRAY = RGBColor(31, 41, 55)    # #1F2937 - Tiêu đề cấp 3 & bold label
COLOR_BODY = RGBColor(17, 24, 39)         # #111827 - Văn bản nội dung
COLOR_MUTED = RGBColor(75, 85, 99)        # #4B5563 - Chú thích phụ

def setup_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    footer = sec.footer
    p_ft = footer.paragraphs[0]
    p_ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_ft.add_run("Trang ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(10)
    r1.font.color.rgb = COLOR_MUTED
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    p_ft._p.append(fldSimple)

def add_styled_heading(doc, text, level):
    style_name = f"Heading {level}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_NAVY
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_DARK_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_DARK_GRAY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    elif level == 4:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_DARK_GRAY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_body_p(doc, text, bold_prefix=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(13)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK_GRAY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_BODY
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r_bullet = p.add_run("•  ")
    r_bullet.font.name = "Times New Roman"
    r_bullet.font.size = Pt(13)
    r_bullet.font.bold = True
    r_bullet.font.color.rgb = COLOR_DARK_BLUE
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(13)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK_GRAY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_BODY
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="0F2C59"/></w:pBdr>')
    pPr.append(shd)
    pPr.append(pBdr)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK_GRAY
    return p

def add_figure_with_notes(doc, image_path, caption_text, observations, explanation=None, ml_implication=None, max_width_inches=5.6):
    abs_img = os.path.join(BASE_DIR, image_path)
    if not os.path.exists(abs_img):
        print(f"[CẢNH BÁO] Không tìm thấy ảnh: {abs_img}")
        p_err = doc.add_paragraph(f"[TODO FIGURE: {image_path}]")
        p_err.runs[0].font.color.rgb = RGBColor(220, 38, 38)
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.paragraph_format.keep_with_next = True
    run_img = p_img.add_run()
    run_img.add_picture(abs_img, width=Inches(max_width_inches))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(6)
    p_cap.paragraph_format.keep_with_next = True
    r_cap = p_cap.add_run(caption_text)
    r_cap.font.name = "Times New Roman"
    r_cap.font.size = Pt(11)
    r_cap.font.bold = True
    r_cap.font.italic = True
    r_cap.font.color.rgb = COLOR_DARK_GRAY
    
    p_obs_title = doc.add_paragraph()
    p_obs_title.paragraph_format.space_before = Pt(2)
    p_obs_title.paragraph_format.space_after = Pt(2)
    p_obs_title.paragraph_format.keep_with_next = True
    r_ot = p_obs_title.add_run("Nhận xét và phân tích biểu đồ:")
    r_ot.font.name = "Times New Roman"
    r_ot.font.size = Pt(12)
    r_ot.font.bold = True
    r_ot.font.color.rgb = COLOR_DARK_BLUE
    
    for item in observations:
        add_bullet_p(doc, item, bold_prefix="Nhận xét: " if item == observations[0] else None)
    if explanation:
        add_bullet_p(doc, explanation, bold_prefix="Giải thích: ")
    if ml_implication:
        add_bullet_p(doc, ml_implication, bold_prefix="Ý nghĩa đối với mô hình / ML: ")
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

def add_styled_table(doc, caption_text, headers, data, col_widths=None, align_cols=None):
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(8)
    p_cap.paragraph_format.space_after = Pt(4)
    p_cap.paragraph_format.keep_with_next = True
    r_cap = p_cap.add_run(caption_text)
    r_cap.font.name = "Times New Roman"
    r_cap.font.size = Pt(11.5)
    r_cap.font.bold = True
    r_cap.font.color.rgb = COLOR_NAVY
    
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_row = table.rows[0]
    hdr_tr = hdr_row._tr.get_or_add_trPr()
    hdr_tr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    hdr_tr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    
    for i, h_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h_text
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F2C59"/>')
        cell._tc.get_or_add_tcPr().append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="0" w:color="0F2C59"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="0F2C59"/><w:left w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        cell._tc.get_or_add_tcPr().append(tcBorders)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        r_tr = row._tr.get_or_add_trPr()
        r_tr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            cell._tc.get_or_add_tcPr().append(shd)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:left w:val="none"/><w:right w:val="none"/></w:tcBorders>')
            cell._tc.get_or_add_tcPr().append(tcBorders)
            p = cell.paragraphs[0]
            if align_cols and c_idx < len(align_cols):
                p.alignment = align_cols[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(10.5)
                r.font.color.rgb = COLOR_BODY
                
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_snippet_with_notes(doc, code_text, caption_text, description_items, source_file=None):
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(8)
    p_cap.paragraph_format.space_after = Pt(2)
    p_cap.paragraph_format.keep_with_next = True
    r_cap = p_cap.add_run(caption_text)
    r_cap.font.name = "Times New Roman"
    r_cap.font.size = Pt(11)
    r_cap.font.bold = True
    r_cap.font.italic = True
    r_cap.font.color.rgb = COLOR_DARK_GRAY
    
    add_code_block(doc, code_text)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(2)
    p_desc.paragraph_format.space_after = Pt(2)
    r_dt = p_desc.add_run("Mô tả:")
    r_dt.font.name = "Times New Roman"
    r_dt.font.size = Pt(11.5)
    r_dt.font.bold = True
    r_dt.font.italic = True
    
    for item in description_items:
        add_bullet_p(doc, item)
        
    if source_file:
        p_src = doc.add_paragraph()
        p_src.paragraph_format.left_indent = Inches(0.25)
        p_src.paragraph_format.space_before = Pt(0)
        p_src.paragraph_format.space_after = Pt(4)
        r_src_label = p_src.add_run("Nguồn: ")
        r_src_label.font.name = "Times New Roman"
        r_src_label.font.size = Pt(11)
        r_src_label.font.italic = True
        
        r_src_code = p_src.add_run(source_file)
        r_src_code.font.name = "Consolas"
        r_src_code.font.size = Pt(10)
        r_src_code.font.color.rgb = COLOR_DARK_GRAY
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
