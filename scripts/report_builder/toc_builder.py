# -*- coding: utf-8 -*-
"""
TOC Builder for Baocao.docx
Builds both the dynamic Word TOC field and pre-rendered TOC paragraphs with dot leaders
so the Table of Contents renders properly in Word and in LibreOffice/PDF.
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .config import add_styled_heading, add_body_p, COLOR_MUTED, COLOR_DARK_BLUE, COLOR_DARK_GRAY, COLOR_BODY

def build_toc(doc, heading_entries=None):
    add_styled_heading(doc, "MỤC LỤC", 1)
    
    p_toc_hint = add_body_p(
        doc,
        "(Mục lục tự động: Trong Microsoft Word, quý độc giả có thể cập nhật bảng mục lục bất kỳ lúc nào bằng cách nhấp chuột phải và chọn 'Update Table' → 'Update entire table')."
    )
    p_toc_hint.runs[0].font.italic = True
    p_toc_hint.runs[0].font.size = Pt(10.5)
    p_toc_hint.runs[0].font.color.rgb = COLOR_MUTED
    p_toc_hint.paragraph_format.space_after = Pt(8)

    # Add dynamic TOC field (for Word native update)
    p_fld = doc.add_paragraph()
    r_fld = p_fld.add_run()
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'separate')
    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'end')
    r_fld._r.append(fld1)
    r_fld._r.append(instr)
    r_fld._r.append(fld2)
    r_fld._r.append(fld3)
    p_fld.paragraph_format.space_after = Pt(0)

    # If pre-rendered entries with page numbers are provided, add them:
    if heading_entries:
        for lvl, title, pno in heading_entries:
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.paragraph_format.line_spacing = 1.15
            
            if lvl == 1:
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
                r_t = p.add_run(title)
                r_t.font.name = "Times New Roman"
                r_t.font.size = Pt(11.5)
                r_t.font.bold = True
                r_t.font.color.rgb = COLOR_DARK_BLUE
            elif lvl == 2:
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                r_t = p.add_run(title)
                r_t.font.name = "Times New Roman"
                r_t.font.size = Pt(10.5)
                r_t.font.bold = True
                r_t.font.color.rgb = COLOR_DARK_GRAY
            else: # lvl == 3
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                r_t = p.add_run(title)
                r_t.font.name = "Times New Roman"
                r_t.font.size = Pt(10)
                r_t.font.color.rgb = COLOR_BODY

            p.add_run("\t")
            r_p = p.add_run(str(pno))
            r_p.font.name = "Times New Roman"
            r_p.font.size = Pt(10.5 if lvl <= 2 else 10)
            if lvl == 1:
                r_p.font.bold = True
                r_p.font.color.rgb = COLOR_DARK_BLUE
            else:
                r_p.font.color.rgb = COLOR_DARK_GRAY

    doc.add_page_break()
