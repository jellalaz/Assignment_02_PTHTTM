#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Master Report Generator Script — Assignment 02
Posts and Telecommunications Institute of Technology (PTIT)
Intelligent System Development

Preserves 100% of the original cover page from report/Baocao_backup.docx.
Generates comprehensive report with 11 chapters + Conclusion matching DungVT's structure.
Uses 100% verified real experimental data from the project.
Includes automated Table of Contents page number calibration.
"""

import os
import sys
import subprocess
import docx
import PyPDF2

# Add scripts directory to sys.path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from report_builder.config import (
    BASE_DIR, REPORT_DOCX, BACKUP_DOCX,
    setup_footer
)
from report_builder.toc_builder import build_toc
from report_builder.ch1_intro import build_chapter_1
from report_builder.ch2_theory import build_chapter_2
from report_builder.ch3_preprocessing import build_chapter_3
from report_builder.ch4_metrics import build_chapter_4
from report_builder.ch5_models import build_chapter_5
from report_builder.ch6_deployment_methods import build_chapter_6
from report_builder.ch7_diabetes import build_chapter_7
from report_builder.ch8_house import build_chapter_8
from report_builder.ch9_ecommerce import build_chapter_9
from report_builder.ch10_comparison import build_chapter_10
from report_builder.ch11_neo4j_conclusion import build_chapter_11

def build_document(heading_entries=None):
    doc = docx.Document(BACKUP_DOCX)

    # 1. Bảo toàn 100% trang bìa gốc (paragraphs 0-14), xóa các paragraph rỗng thừa phía sau
    body = doc._body._body
    for p in list(doc.paragraphs)[15:]:
        body.remove(p._p)

    # 2. Thiết lập Footer trang (Trang 2 trở đi có số trang)
    setup_footer(doc)

    # 3. Ngắt trang sau bìa để vào Mục lục
    doc.add_page_break()

    # 4. Tạo Mục lục (gồm cả dynamic TOC field lẫn bảng mục lục chi tiết có số trang)
    build_toc(doc, heading_entries)

    # 5. Xây dựng lần lượt 11 chương
    build_chapter_1(doc)
    build_chapter_2(doc)
    build_chapter_3(doc)
    build_chapter_4(doc)
    build_chapter_5(doc)
    build_chapter_6(doc)
    build_chapter_7(doc)
    build_chapter_8(doc)
    build_chapter_9(doc)
    build_chapter_10(doc)
    build_chapter_11(doc)

    doc.save(REPORT_DOCX)
    return doc

def extract_heading_pages_from_pdf(pdf_path, headings_list):
    reader = PyPDF2.PdfReader(pdf_path)
    total_pages = len(reader.pages)
    page_texts = [reader.pages[i].extract_text() for i in range(total_pages)]
    
    # Identify where content starts (search for Chapter I title after page 2)
    start_content_page = 4  # Default index 4 = Page 5
    for p_idx in range(1, min(10, total_pages)):
        if "CHƯƠNG I. GIỚI THIỆU BÀI TOÁN" in page_texts[p_idx] and p_idx > 1:
            # Check if this page is not a TOC page (TOC pages contain "MỤC LỤC" or multiple chapters)
            if "MỤC LỤC" not in page_texts[p_idx] and "CHƯƠNG II." not in page_texts[p_idx]:
                start_content_page = p_idx
                break
            
    heading_pages = []
    current_search_page = start_content_page
    
    for lvl, title in headings_list:
        found_page = None
        # Clean title snippet for searching
        snippet = title.replace("—", "").replace("-", "").strip()[:30]
        for p_idx in range(current_search_page, total_pages):
            if snippet in page_texts[p_idx]:
                found_page = p_idx + 1  # 1-indexed
                current_search_page = p_idx
                break
        if not found_page:
            found_page = current_search_page + 1
        heading_pages.append((lvl, title, found_page))
        
    return heading_pages

def collect_headings_from_doc(doc):
    headings = []
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and p.text.strip() != 'MỤC LỤC':
            lvl = int(p.style.name.split()[1])
            if lvl <= 3:  # Only include levels 1, 2, 3 in TOC
                headings.append((lvl, p.text.strip()))
    return headings

def main():
    print("=" * 70)
    print("BẮT ĐẦU HOÀN THIỆN BÁO CÁO CHÍNH THỨC Baocao.docx (PHIÊN BẢN V2 MỞ RỘNG)")
    print("=" * 70)

    if not os.path.exists(BACKUP_DOCX):
        print(f"[LỖI] Không tìm thấy file backup trang bìa: {BACKUP_DOCX}")
        sys.exit(1)

    pdf_path = os.path.join(BASE_DIR, "report", "Baocao.pdf")

    # Step 1: Check if an existing PDF is already available to get initial TOC estimates
    print("\n>>> [BƯỚC 1/3] Tạo bản thảo có danh mục cấu trúc...")
    doc_initial = build_document(heading_entries=None)
    headings = collect_headings_from_doc(doc_initial)
    print(f"    Đã xác định {len(headings)} đề mục Heading (1-3).")

    # Step 2: Build intermediate document with placeholder pages and export to PDF
    dummy_entries = [(lvl, title, 5 + i // 2) for i, (lvl, title) in enumerate(headings)]
    print("\n>>> [BƯỚC 2/3] Xây dựng văn bản với cấu trúc bảng Mục lục 3 trang...")
    build_document(heading_entries=dummy_entries)
    print("    Đang xuất bản PDF trung gian qua LibreOffice...")
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", REPORT_DOCX, "--outdir", os.path.join(BASE_DIR, "report")],
        check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )

    # Step 3: Extract exact heading pages from PDF (after TOC pages)
    print("\n>>> [BƯỚC 3/3] Trích xuất số trang chính xác và hoàn thiện tài liệu...")
    exact_pages = extract_heading_pages_from_pdf(pdf_path, headings)
    build_document(heading_entries=exact_pages)
    
    # Final PDF export
    print("    Đang xuất bản file PDF chính thức cuối cùng...")
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", REPORT_DOCX, "--outdir", os.path.join(BASE_DIR, "report")],
        check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )

    reader = PyPDF2.PdfReader(pdf_path)
    total_pages = len(reader.pages)
    print("=" * 70)
    print("XUẤT BẢN THÀNH CÔNG BÁO CÁO CHÍNH THỨC V2!")
    print(f"• File Word: {REPORT_DOCX}")
    print(f"• File PDF:  {pdf_path} ({total_pages} trang)")
    print("=" * 70)

if __name__ == "__main__":
    main()
