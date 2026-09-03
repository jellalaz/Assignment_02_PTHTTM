import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_document(filepath):
    doc = docx.Document(filepath)
    
    # Track fixes for logging
    fixes_made = 0
    
    for p in doc.paragraphs:
        # Style formatting
        if p.style.name.startswith('Heading'):
            lvl = p.style.name.split()[-1]
            if lvl == '1':
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.font.bold = True
            elif lvl == '2':
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True
            elif lvl == '3':
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.bold = True
        else:
            # Body text formatting
            # p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # p.paragraph_format.line_spacing = 1.3
            for run in p.runs:
                if not run.font.name:
                    run.font.name = 'Times New Roman'
                if not run.font.size:
                    run.font.size = Pt(13)
                    
        # Text replacement
        if 'Ridge Regression' in p.text:
            text = p.text
            new_text = text.replace('Ridge Regression', 'Gradient Boosting Regressor')
            # Careful replacement preserving runs if possible, but simplest is clearing runs and adding new one
            # if we just need content replaced.
            p.text = new_text
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
            fixes_made += 1
            
        if 'Ridge' in p.text and 'Regression' not in p.text:
             # Just in case there's isolated 'Ridge'
             p.text = p.text.replace('Ridge', 'Gradient Boosting')
             for run in p.runs:
                 run.font.name = 'Times New Roman'
                 run.font.size = Pt(13)
             fixes_made += 1

    # Fix tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                    if 'Ridge Regression' in p.text:
                        p.text = p.text.replace('Ridge Regression', 'Gradient Boosting Regressor')
                        fixes_made += 1
                    elif 'Ridge' in p.text:
                        p.text = p.text.replace('Ridge', 'Gradient Boosting')
                        fixes_made += 1

    doc.save(filepath)
    return fixes_made

if __name__ == '__main__':
    doc_path = '/home/jellalaz/Documents/Jellalaz/DATA_CODE/PYTHON/Assignment_02/report/Baocao.docx'
    fixes = fix_document(doc_path)
    print(f"Made {fixes} replacements of Ridge -> Gradient Boosting")
